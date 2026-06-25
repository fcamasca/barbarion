from __future__ import annotations

from pathlib import Path, PurePosixPath

import pytest
from docx import Document

from barbarion.domain.models import (
    Confidence,
    DiscoveredFile,
    ExtractionContext,
    SourceFile,
)
from barbarion.infrastructure.parsers.docx import DocxParser
from barbarion.infrastructure.parsers.encoding import (
    DOCUMENT_CORRUPT,
    EXTRACTION_LIMIT_EXCEEDED,
    TextExtractionError,
)


def source_for(path: Path, root: Path | None = None) -> SourceFile:
    effective_root = path.parent if root is None else root
    stat_result = path.stat()
    return SourceFile(
        discovered=DiscoveredFile(
            root=effective_root,
            relative_path=PurePosixPath(path.relative_to(effective_root).as_posix()),
            runtime_path=path,
            extension=path.suffix,
            size_bytes=stat_result.st_size,
            mtime_ns=stat_result.st_mtime_ns,
        )
    )


def context(max_extracted_chars: int = 10000) -> ExtractionContext:
    return ExtractionContext(
        encodings=("utf-8", "cp1252", "latin-1"),
        max_extracted_chars=max_extracted_chars,
        max_pdf_pages=10,
    )


def write_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Guia", level=1)
    document.add_paragraph("Parrafo inicial")
    document.add_heading("Detalle", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Campo"
    table.cell(0, 1).text = "Valor"
    table.cell(1, 0).text = "Estado"
    table.cell(1, 1).text = "Activo"
    document.add_paragraph("Cierre")
    document.save(path)


def test_docx_parser_extracts_headings_paragraphs_and_tables_in_order(
    tmp_path: Path,
) -> None:
    path = tmp_path / "manual.docx"
    write_docx(path)

    result = DocxParser().extract(source_for(path), context())

    assert result.text == (
        "Guia\n\n"
        "Parrafo inicial\n\n"
        "Detalle\n\n"
        "Campo | Valor\nEstado | Activo\n\n"
        "Cierre"
    )
    assert [unit.unit_type for unit in result.units] == [
        "heading",
        "paragraph",
        "heading",
        "table",
        "paragraph",
    ]
    assert [unit.metadata["ordinal"] for unit in result.units] == [0, 1, 2, 3, 4]
    assert result.units[1].metadata["breadcrumb"] == ("Guia",)
    assert result.units[3].metadata["breadcrumb"] == ("Guia", "Detalle")
    assert all(unit.confidence == Confidence.HIGH for unit in result.units)


def test_docx_parser_applies_text_limit(tmp_path: Path) -> None:
    path = tmp_path / "manual.docx"
    write_docx(path)

    with pytest.raises(TextExtractionError) as raised:
        DocxParser().extract(source_for(path), context(max_extracted_chars=10))

    assert raised.value.to_pipeline_error().error_code == EXTRACTION_LIMIT_EXCEEDED


def test_docx_parser_reports_corrupt_document(tmp_path: Path) -> None:
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a zip")

    with pytest.raises(TextExtractionError) as raised:
        DocxParser().extract(source_for(path), context())

    assert raised.value.to_pipeline_error().error_code == DOCUMENT_CORRUPT
