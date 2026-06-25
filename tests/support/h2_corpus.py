"""Construccion de corpus sintetico para pruebas H2."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from pypdf import PdfWriter
from reportlab.pdfgen import canvas


FIXTURE_ROOT = Path(__file__).resolve().parents[1] / "fixtures" / "h2_corpus"

TEXT_EXTENSIONS = {
    ".sql",
    ".pks",
    ".pkb",
    ".prc",
    ".fnc",
    ".trg",
    ".pck",
    ".vw",
    ".vws",
    ".pkg",
    ".tps",
    ".srw",
    ".sru",
    ".srf",
    ".srm",
    ".srj",
    ".srd",
    ".md",
    ".txt",
    ".yaml",
    ".yml",
    ".json",
    ".ini",
}
GENERATED_EXTENSIONS = {".pdf", ".docx"}
EXPECTED_EXTENSIONS = TEXT_EXTENSIONS | GENERATED_EXTENSIONS | {".pbl"}


def build_h2_corpus(target: Path, *, include_errors: bool = False) -> Path:
    """Copia fixtures textuales y genera documentos binarios sinteticos."""
    if target.exists():
        shutil.rmtree(target)
    shutil.copytree(FIXTURE_ROOT, target)
    _write_binary_samples(target)
    _write_cp1252_sample(target)
    _write_large_text(target)
    if not include_errors:
        shutil.rmtree(target / "errors")
    else:
        _write_error_samples(target)
    return target


def fixture_inventory(root: Path) -> dict[str, int]:
    """Cuenta extensiones presentes en un corpus ya construido."""
    inventory: dict[str, int] = {}
    for path in root.rglob("*"):
        if path.is_file():
            inventory[path.suffix.lower()] = inventory.get(path.suffix.lower(), 0) + 1
    return inventory


def _write_binary_samples(root: Path) -> None:
    docs = root / "docs"
    docs.mkdir(parents=True, exist_ok=True)
    _write_text_pdf(docs / "manual.pdf", ("Pagina uno del manual", "Pagina dos"))
    _write_docx(docs / "manual.docx")
    (root / "powerbuilder" / "binary_library.pbl").write_bytes(
        b"\x00PBL\x00synthetic\xff\x10"
    )


def _write_error_samples(root: Path) -> None:
    errors = root / "errors"
    errors.mkdir(parents=True, exist_ok=True)
    (errors / "invalid_utf8.txt").write_bytes(b"texto valido\n\xff\xfe\xfa")
    (errors / "corrupt.pdf").write_bytes(b"not a real pdf")
    (errors / "corrupt.docx").write_bytes(b"not a real docx")
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with (errors / "blank.pdf").open("wb") as output:
        writer.write(output)


def _write_cp1252_sample(root: Path) -> None:
    content = "$PBExportHeader$uo_cp1252.sru\n" "global type uo_cp1252 from userobject\n" "end type\n" "public function string nombre();\n" 'return "Acento: ñ"\n' "end function\n"
    (root / "powerbuilder" / "cp1252_export.sru").write_bytes(
        content.encode("cp1252")
    )


def _write_large_text(root: Path) -> None:
    lines = [f"Linea sintetica {index:03d}" for index in range(120)]
    (root / "docs" / "large.txt").write_text("\n\n".join(lines), encoding="utf-8")


def _write_text_pdf(path: Path, pages: tuple[str, ...]) -> None:
    pdf_canvas = canvas.Canvas(str(path))
    for page_text in pages:
        pdf_canvas.drawString(72, 720, page_text)
        pdf_canvas.showPage()
    pdf_canvas.save()


def _write_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Manual sintetico", level=1)
    document.add_paragraph("Documento generado para pruebas de ingesta.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "clave"
    table.cell(0, 1).text = "valor"
    table.cell(1, 0).text = "modo"
    table.cell(1, 1).text = "local"
    document.save(path)
